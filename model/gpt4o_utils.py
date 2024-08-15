import pdfplumber
import json
from openai import OpenAI
from dotenv import load_dotenv
import os
from docx import Document

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    return text

# detect section titles using GPT-4
def detect_section_titles(text):

    prompt = f"Extract section titles from the following text:\n\n{text}"  

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an assistant that extracts section titles from academic papers."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=300,
        n=1,
        stop=None,
        temperature=0.7)

    section_titles = response.choices[0].message.content.strip().split('\n')
    return [title.strip() for title in section_titles if title.strip()]

def split_text_into_chunks(text, max_length=3000):
    words = text.split()
    chunks = []
    current_chunk = []

    for word in words:
        if len(''.join(current_chunk)) + len(word) + 1 > max_length:
            chunks.append(' '.join(current_chunk))
            current_chunk = []
        current_chunk.append(word)
    if current_chunk:
        chunks.append(' '.join(current_chunk)) 
    return chunks

def extract_json_content(response_content):
    try:
        json_start = response_content.index("{")
        json_end = response_content.rindex("}") + 1
        json_content = response_content[json_start:json_end]
        return json.loads(json_content)
    except (ValueError, json.JSONDecodeError):
        print(f"Failed to extract JSON content from response: {response_content}")
        return None

# separate sections based on detected titles
def separate_sections(text, section_titles):
    chunks = split_text_into_chunks(text, max_length=5000)  
    all_section_texts = {title: "" for title in section_titles}

    for chunk in chunks:
        prompt = f"""
        Pleae separate the following text into sections. Return the result as a JSON dictionary where the keys are the section titles and the values are the corresponding content.

        Section Titles:
        {section_titles}

        Text:
        {chunk}
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2048,
            n=1,
            temperature=0.7
        )

        response_content = response.choices[0].message.content.strip()
        section_texts = extract_json_content(response_content)
        if section_texts:
            for title, content in section_texts.items():
                if isinstance(content, str):
                    if title in all_section_texts:
                        all_section_texts[title] += content + "\n"
                    else:
                        all_section_texts[title] = content + "\n"
                elif isinstance(content, dict):
                    if title in all_section_texts:
                        all_section_texts[title] += json.dumps(content) + "\n"
                    else:
                        all_section_texts[title] = json.dumps(content) + "\n"
    return all_section_texts

# generate comments using GPT-4
def generate_comment(section, content, criteria):
    prompt = f"""
    You are a peer reviewer for a scientific paper. Please provide detailed and constructive feedback on the {section} section of the paper, in close accordance to the provided criteria. 
    Please ensure your comments are professional, thorough, and helpful for the authors to improve their work.

    Section: {section}
    
    Content: {content}
    
    Criteria: {criteria}
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000,
        n=1,
        temperature=0.7
    )

    return response.choices[0].message.content.strip()

def save_to_json(sections, filename):
    with open(filename, 'w') as json_file:
        json.dump(sections, json_file, indent=4)

def load_from_json(filename):
    with open(filename, 'r') as json_file:
        return json.load(json_file)

def generate_comments_for_sections(sections, criteria):
    comments = {}
    for section, content in sections.items():
        comments[section] = generate_comment(section, content, criteria)
    entire_paper_content = "\n".join(sections.values())
    comments['Novelty and Merit'] = check_novelty_of_paper(entire_paper_content)
    return comments

def save_comments_to_json(comments,filename):
    with open(filename, 'w') as json_file:
        json.dump(comments, json_file, indent=4)   

def load_journal_criteria(filename):
    text = extract_text_from_pdf(filename)
    criteria = text.split('\n')
    return {'criteria': criteria}

def revise_paper(sections, criteria):
    revised_sections = {}
    for section, content in sections.items():
        prompt = f"""
        The following is the content of the {section} section of an academic paper. Revise it according to the provided criteria.
        
        Section Content:
        {content}
        
        Criteria:
        {criteria}
        """

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2048,
            n=1,
            temperature=0.7
        )

        revised_sections[section] = response.choices[0].message.content.strip()
    return revised_sections

def check_novelty_of_paper(content):
    prompt = f"""
    Please evaluate the novelty of the paper proposal based on prior literature and provide a score from 1 to 5 on the novelty and merit of the paper, with 1 being the lowest and 5 being the highest. 
    Justify your score with a brief explanation.
    
    Content: {content}
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=400,
        n=1,
        temperature=0.7
    )

    return response.choices[0].message.content.strip()

def export_revised_paper_to_json(revised_sections, filename):
    with open(filename, 'w') as json_file:
        json.dump(revised_sections, json_file, indent=4)
        
def export_json_to_docx(json_file, filename):
    doc = Document()
    for section, content in json_file.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    doc.save(filename)
    print(f"documents saved as {filename}")
    
# main function
def main():
    input_type = input("Enter '1' for paper proposal, '2' for paper, '3' for patent application: ")
    paper_path = "../data/sample_patents/patent example.pdf"
    # paper_path = "../data/sample_proposals/NSF-Proposal-Example-2010.pdf"
    criteria_path = "../data/criteria/patent_criteria.pdf"
    # criteria_path = "../data/criteria/nsf_criteria.pdf"
    paper_file_name = os.path.splitext(os.path.basename(paper_path))[0]

    paper_text = extract_text_from_pdf(paper_path)
    criteria_json = load_journal_criteria(criteria_path)

    section_titles = detect_section_titles(paper_text)
    section_texts = separate_sections(paper_text, section_titles)
    sections_filename = f"{paper_file_name}_sections.json"
    save_to_json(section_texts, sections_filename)

    sections = load_from_json(sections_filename) 
    
    if input_type == '1' or input_type == '3':
        comments = generate_comments_for_sections(sections, criteria_json['criteria'])
        comments_filename = f"{paper_file_name}_comments.json"
        save_comments_to_json(comments, comments_filename)
        print(f"comments saved to {comments_filename}")
        if input("Do you want to export the comments to a document? (yes/no)") == 'yes':
            doc_filename = f"{paper_file_name}_comments.docx"
            export_json_to_docx(comments, doc_filename)
        
    elif input_type == '2':
        user_choice = input("Enter '1' to generate comments, '2' to revise:")
        if user_choice == '1':
            comments = generate_comments_for_sections(sections, criteria_json['criteria'])
            comments_filename = f"{paper_file_name}_comments.json"
            save_comments_to_json(comments, comments_filename)
            print(f"comments saved to {comments_filename}")
            if input("Do you want to export the comments to a document? (yes/no)") == 'yes':
                doc_filename = f"{paper_file_name}_comments.docx"
                export_json_to_docx(comments, doc_filename)
        
        elif user_choice == '2':
            revised_sections = revise_paper(sections, criteria_json['criteria'])
            revised_filename = f"{paper_file_name}_revised.json"
            export_revised_paper_to_json(revised_sections, revised_filename)
            print(f"revised paper exported as {revised_filename}")
            if input("Do you want to export the revised paper to a document? (yes/no)") == 'yes':
                doc_filename = f"{paper_file_name}_revised.docx"
                export_json_to_docx(revised_sections, doc_filename)

if __name__ == "__main__":
    main()
